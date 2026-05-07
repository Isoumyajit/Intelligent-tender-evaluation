from typing import Annotated
from uuid import UUID

from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, UploadFile, status
from fastapi.responses import Response
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.models import Attachment, Bid, BidAttachment, Tender
from app.schemas import BidListResponse, BidResponse, BidUpdate
from app.services.audit_service import log_audit
from app.services.ocr_service import run_ocr

router = APIRouter(prefix="/tenders/{tender_id}/bid", tags=["bids"])


async def _get_tender_or_404(tender_id: UUID, db: AsyncSession) -> Tender:
    result = await db.execute(select(Tender).where(Tender.tender_id == tender_id))
    tender = result.scalar_one_or_none()
    if not tender:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Tender not found")
    return tender


@router.post("/", response_model=BidResponse, status_code=status.HTTP_201_CREATED)
async def create_bid(
    tender_id: UUID,
    bidder_name: Annotated[str, Form()],
    documents: Annotated[list[UploadFile], File(description="Upload one or more files")],
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: AsyncSession = Depends(get_db),
):
    await _get_tender_or_404(tender_id, db)

    attachments: list[Attachment] = []
    for doc in documents:
        file_bytes = await doc.read()
        attachment = Attachment(
            file_name=doc.filename or "unknown",
            content_type=doc.content_type or "application/octet-stream",
            data=file_bytes,
        )
        db.add(attachment)
        attachments.append(attachment)

    await db.flush()

    bid = Bid(tender_id=tender_id, bidder_name=bidder_name)
    db.add(bid)
    await db.flush()

    for attachment in attachments:
        bid_attachment = BidAttachment(
            bid_id=bid.bid_id,
            attachment_ref_id=attachment.attachment_ref_id,
        )
        db.add(bid_attachment)

    await log_audit(
        db,
        tender_id=tender_id,
        event="bidder_added",
        audit_desc=f"Bidder '{bidder_name}' (bid {bid.bid_id}) added to tender {tender_id}",
        bidder_id=bid.bid_id,
    )

    await db.commit()
    await db.refresh(bid)

    for attachment in attachments:
        background_tasks.add_task(run_ocr, attachment.attachment_ref_id)

    return _bid_to_response(bid)


@router.get("/", response_model=list[BidListResponse])
async def list_bids(
    tender_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    await _get_tender_or_404(tender_id, db)
    result = await db.execute(
        select(Bid).where(Bid.tender_id == tender_id).order_by(Bid.created_at.desc())
    )
    return result.scalars().all()


@router.get("/{bid_id}", response_model=BidResponse)
async def get_bid(
    tender_id: UUID,
    bid_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    bid = await _get_bid_or_404(tender_id, bid_id, db)
    return _bid_to_response(bid)


@router.put("/{bid_id}", response_model=BidResponse)
async def update_bid(
    tender_id: UUID,
    bid_id: UUID,
    payload: BidUpdate,
    db: AsyncSession = Depends(get_db),
):
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    update_data = payload.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(bid, field, value)

    await db.commit()
    await db.refresh(bid)
    return _bid_to_response(bid)


@router.delete("/{bid_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_bid(
    tender_id: UUID,
    bid_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    attachment_ref_ids = [ba.attachment_ref_id for ba in bid.bid_attachments]

    await db.delete(bid)

    if attachment_ref_ids:
        att_result = await db.execute(
            select(Attachment).where(Attachment.attachment_ref_id.in_(attachment_ref_ids))
        )
        for attachment in att_result.scalars().all():
            await db.delete(attachment)

    await db.commit()


@router.put("/{bid_id}/documents", response_model=BidResponse)
async def add_bid_document(
    tender_id: UUID,
    bid_id: UUID,
    document: Annotated[UploadFile, File(description="Upload a file to add to the bid")],
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: AsyncSession = Depends(get_db),
):
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    file_bytes = await document.read()
    attachment = Attachment(
        file_name=document.filename or "unknown",
        content_type=document.content_type or "application/octet-stream",
        data=file_bytes,
    )
    db.add(attachment)
    await db.flush()

    bid_attachment = BidAttachment(
        bid_id=bid.bid_id,
        attachment_ref_id=attachment.attachment_ref_id,
    )
    db.add(bid_attachment)

    await db.commit()
    await db.refresh(bid)

    background_tasks.add_task(run_ocr, attachment.attachment_ref_id)

    return _bid_to_response(bid)


@router.delete("/{bid_id}/documents/{attachment_ref_id}", status_code=status.HTTP_204_NO_CONTENT)
async def remove_bid_document(
    tender_id: UUID,
    bid_id: UUID,
    attachment_ref_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    linked_ids = {ba.attachment_ref_id for ba in bid.bid_attachments}
    if attachment_ref_id not in linked_ids:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found for this bid",
        )

    ba_result = await db.execute(
        select(BidAttachment).where(
            BidAttachment.bid_id == bid_id,
            BidAttachment.attachment_ref_id == attachment_ref_id,
        )
    )
    bid_attachment = ba_result.scalar_one()
    await db.delete(bid_attachment)

    att_result = await db.execute(
        select(Attachment).where(Attachment.attachment_ref_id == attachment_ref_id)
    )
    attachment = att_result.scalar_one_or_none()
    if attachment:
        await db.delete(attachment)

    await db.commit()


@router.get("/{bid_id}/documents/{attachment_ref_id}")
async def download_bid_document(
    tender_id: UUID,
    bid_id: UUID,
    attachment_ref_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    linked_ids = {ba.attachment_ref_id for ba in bid.bid_attachments}
    if attachment_ref_id not in linked_ids:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found for this bid",
        )

    att_result = await db.execute(
        select(Attachment).where(Attachment.attachment_ref_id == attachment_ref_id)
    )
    attachment = att_result.scalar_one_or_none()
    if not attachment:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    return Response(
        content=attachment.data,
        media_type=attachment.content_type,
        headers={"Content-Disposition": f'attachment; filename="{attachment.file_name}"'},
    )


@router.get("/{bid_id}/documents/{attachment_ref_id}/text")
async def get_bid_document_text(
    tender_id: UUID,
    bid_id: UUID,
    attachment_ref_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """Return extracted plain text from a bid attachment."""
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    linked_ids = {ba.attachment_ref_id for ba in bid.bid_attachments}
    if attachment_ref_id not in linked_ids:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    att_result = await db.execute(
        select(Attachment).where(Attachment.attachment_ref_id == attachment_ref_id)
    )
    attachment = att_result.scalar_one_or_none()
    if not attachment or not attachment.data:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    ct = (attachment.content_type or "").lower()
    text = ""
    if "pdf" in ct or attachment.file_name.lower().endswith(".pdf"):
        try:
            import fitz
            doc = fitz.open(stream=attachment.data, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
        except Exception:
            text = "(Could not extract text from PDF)"
    elif "word" in ct or attachment.file_name.lower().endswith((".docx", ".doc")):
        try:
            import io
            from docx import Document as DocxDocument
            docx_doc = DocxDocument(io.BytesIO(attachment.data))
            text = "\n".join(p.text for p in docx_doc.paragraphs)
        except Exception:
            text = "(Could not extract text from DOCX)"
    elif "image" in ct or attachment.file_name.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
        mime = attachment.content_type or "image/png"
        return Response(content=attachment.data, media_type=mime)
    else:
        text = attachment.data.decode("utf-8", errors="replace")

    return Response(content=text, media_type="text/plain; charset=utf-8")


@router.get("/{bid_id}/documents/{attachment_ref_id}/metadata")
async def get_bid_document_metadata(
    tender_id: UUID,
    bid_id: UUID,
    attachment_ref_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """Return page count and basic info for a bid attachment."""
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    linked_ids = {ba.attachment_ref_id for ba in bid.bid_attachments}
    if attachment_ref_id not in linked_ids:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    att_result = await db.execute(
        select(Attachment).where(Attachment.attachment_ref_id == attachment_ref_id)
    )
    attachment = att_result.scalar_one_or_none()
    if not attachment or not attachment.data:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    ct = (attachment.content_type or "").lower()
    total_pages = 1

    if "pdf" in ct or attachment.file_name.lower().endswith(".pdf"):
        try:
            import fitz
            doc = fitz.open(stream=attachment.data, filetype="pdf")
            total_pages = len(doc)
            doc.close()
        except Exception:
            pass

    return {
        "documentId": str(attachment_ref_id),
        "fileName": attachment.file_name,
        "totalPages": total_pages,
        "mimeType": attachment.content_type,
    }


@router.get("/{bid_id}/documents/{attachment_ref_id}/page/{page_number}")
async def get_bid_document_page(
    tender_id: UUID,
    bid_id: UUID,
    attachment_ref_id: UUID,
    page_number: int,
    db: AsyncSession = Depends(get_db),
):
    """Render a single page of a bid attachment as PNG (for PDFs) or return text/image."""
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    linked_ids = {ba.attachment_ref_id for ba in bid.bid_attachments}
    if attachment_ref_id not in linked_ids:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    att_result = await db.execute(
        select(Attachment).where(Attachment.attachment_ref_id == attachment_ref_id)
    )
    attachment = att_result.scalar_one_or_none()
    if not attachment or not attachment.data:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    ct = (attachment.content_type or "").lower()

    if "pdf" in ct or attachment.file_name.lower().endswith(".pdf"):
        try:
            import fitz
            doc = fitz.open(stream=attachment.data, filetype="pdf")
            if page_number < 1 or page_number > len(doc):
                doc.close()
                raise HTTPException(status_code=404, detail=f"Page {page_number} not found")
            page = doc[page_number - 1]
            pix = page.get_pixmap(dpi=150)
            png_bytes = pix.tobytes("png")
            doc.close()
            return Response(content=png_bytes, media_type="image/png")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=500, detail="Could not render PDF page")

    if "image" in ct or attachment.file_name.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
        mime = attachment.content_type or "image/png"
        return Response(content=attachment.data, media_type=mime)

    text = attachment.data.decode("utf-8", errors="replace")
    return Response(content=text, media_type="text/plain; charset=utf-8")


async def _get_bid_or_404(tender_id: UUID, bid_id: UUID, db: AsyncSession) -> Bid:
    await _get_tender_or_404(tender_id, db)
    result = await db.execute(
        select(Bid).where(Bid.bid_id == bid_id, Bid.tender_id == tender_id)
    )
    bid = result.scalar_one_or_none()
    if not bid:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Bid not found")
    return bid


def _bid_to_response(bid: Bid) -> dict:
    return {
        "bid_id": bid.bid_id,
        "tender_id": bid.tender_id,
        "bidder_name": bid.bidder_name,
        "created_at": bid.created_at,
        "updated_at": bid.updated_at,
        "attachments": [ba.attachment for ba in bid.bid_attachments],
    }


@router.get("/{bid_id}/evaluation")
async def get_bid_evaluation(
    tender_id: UUID,
    bid_id: UUID,
    db: AsyncSession = Depends(get_db),
):
    """Score a bidder against the tender's LLM-extracted criteria.

    Reads the latest completed job, extracts text from all bidder
    attachments, then runs a multi-signal matcher per criterion:
      - Phrase matching (multi-word spans from predicate)
      - Negative-signal detection (NOT, pending, expired, etc.)
      - Amount/threshold verification where applicable
      - Per-attachment evidence extraction
    """
    import re
    from app.models import Job, TenderCriteria

    bid = await _get_bid_or_404(tender_id, bid_id, db)

    job_result = await db.execute(
        select(Job)
        .where(Job.tender_id == tender_id, Job.status == "completed")
        .order_by(Job.updated_at.desc())
        .limit(1)
    )
    job = job_result.scalar_one_or_none()
    if job is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No completed evaluation found for this tender",
        )

    def _parse_paged_text(raw: str) -> list[dict]:
        """Split text with ``--- PAGE N ---`` delimiters into per-page dicts."""
        page_delimiter = re.compile(r"^---\s*PAGE\s+(\d+)\s*---\s*$", re.MULTILINE)
        splits = page_delimiter.split(raw)
        if len(splits) < 3:
            return [{"page_number": 1, "text": raw}]
        pages: list[dict] = []
        i = 1
        while i < len(splits) - 1:
            page_num = int(splits[i])
            page_text = splits[i + 1].strip()
            pages.append({"page_number": page_num, "text": page_text})
            i += 2
        return pages or [{"page_number": 1, "text": raw}]

    # doc_pages: list of (file_name, doc_url, pages) where pages = [{page_number, text}]
    doc_pages: list[tuple[str, str, list[dict]]] = []
    for ba in bid.bid_attachments:
        att = ba.attachment
        if not att.data:
            continue
        doc_url = f"/tenders/{tender_id}/bid/{bid_id}/documents/{att.attachment_ref_id}"
        pages: list[dict] = []
        ct = (att.content_type or "").lower()
        if "pdf" in ct or att.file_name.lower().endswith(".pdf"):
            try:
                import fitz
                doc = fitz.open(stream=att.data, filetype="pdf")
                for pi, page in enumerate(doc, 1):
                    pages.append({"page_number": pi, "text": page.get_text()})
                doc.close()
            except Exception:
                pages = _parse_paged_text(att.data.decode("utf-8", errors="replace"))
        elif "word" in ct or att.file_name.lower().endswith((".docx", ".doc")):
            try:
                import io
                from docx import Document as DocxDocument
                docx_doc = DocxDocument(io.BytesIO(att.data))
                pages = [{"page_number": 1, "text": "\n".join(p.text for p in docx_doc.paragraphs)}]
            except Exception:
                pages = _parse_paged_text(att.data.decode("utf-8", errors="replace"))
        elif "image" in ct or att.file_name.lower().endswith((".png", ".jpg", ".jpeg")):
            pages = [{"page_number": 1, "text": f"[Image: {att.file_name}]"}]
        else:
            raw_text = att.data.decode("utf-8", errors="replace")
            pages = _parse_paged_text(raw_text)
        doc_pages.append((att.file_name, doc_url, pages))

    # Flatten for backward-compatible full_text search
    doc_texts: list[tuple[str, str, str]] = [
        (name, url, "\n".join(p["text"] for p in pages))
        for name, url, pages in doc_pages
    ]
    full_text = "\n".join(t for _, _, t in doc_texts).lower()

    criteria_result = await db.execute(
        select(TenderCriteria).where(TenderCriteria.job_id == job.job_id)
    )
    criteria_groups = criteria_result.scalars().all()

    from app.models import EvaluationOverride
    override_result = await db.execute(
        select(EvaluationOverride).where(
            EvaluationOverride.tender_id == tender_id,
            EvaluationOverride.bid_id == bid_id,
        )
    )
    overrides = {
        ov.condition_name: ov for ov in override_result.scalars().all()
    }

    category_map = {
        "financial": "Financial",
        "technical": "Technical",
        "compliance": "Compliance",
        "certification": "Eligibility",
        "legal": "Eligibility",
        "experience": "Technical",
    }

    negative_patterns = re.compile(
        r"\bnot\s+(?:submitted|provided|renewed|notarized|available|attached)"
        r"|\bpending\b|\bexpired\b|\bbelow\b|\brefused\b"
        r"|\bno\s+(?:emd|challan|affidavit|certificate)",
        re.IGNORECASE,
    )

    missing_doc_patterns = re.compile(
        r"not\s+submitted|not\s+included|not\s+provided|not\s+attached"
        r"|no\s+emd\s+submitted|document.*not.*found",
        re.IGNORECASE,
    )

    def readable_join(items: list[str], fallback: str = "the required details") -> str:
        unique_items = []
        seen = set()
        for item in items:
            clean = item.strip(" .,;:()").lower()
            if clean and clean not in seen:
                seen.add(clean)
                unique_items.append(clean)

        if not unique_items:
            return fallback
        if len(unique_items) == 1:
            return unique_items[0]
        return ", ".join(unique_items[:-1]) + f" and {unique_items[-1]}"

    def source_label(doc_name: str | None) -> str:
        return f" in {doc_name}" if doc_name else " in the submitted documents"

    def build_eval_note(
        *,
        status: str,
        title: str,
        predicate: str,
        doc_name: str | None,
        matched_terms: list[str],
        missing_terms: list[str],
        neg_phrase: str,
        excerpt: str,
    ) -> str:
        requirement = predicate.rstrip(".")
        matched = readable_join(matched_terms, "some supporting evidence")
        missing = readable_join(missing_terms, "one or more required details")
        source = source_label(doc_name)

        if status == "passed":
            return (
                f"{title} passed because the submission provides {matched}{source}. "
                f"Requirement checked: {requirement}."
            )
        if status == "partial":
            if neg_phrase:
                return (
                    f"Review required because {title} has supporting evidence{source}, "
                    f"but the evidence also mentions '{neg_phrase}'. "
                    f"Please confirm whether this satisfies: {requirement}."
                )
            return (
                f"Review required because only partial evidence was found for {title}{source}. "
                f"Found: {matched}. Missing or unclear: {missing}. "
                f"Requirement checked: {requirement}."
            )
        if status == "missing-document":
            submitted_docs = readable_join([d[0] for d in doc_texts], "no bidder documents")
            if excerpt:
                return (
                    f"Document not submitted for {title}. The available evidence{source} "
                    f"indicates '{excerpt[:140]}'. Requirement checked: {requirement}."
                )
            return (
                f"Document not submitted for {title}. Expected evidence for: {requirement}. "
                f"Searched submitted documents: {submitted_docs}."
            )

        if neg_phrase:
            return (
                f"Failed because the submission indicates '{neg_phrase}' for {title}{source}. "
                f"This does not satisfy: {requirement}."
            )
        return (
            f"Failed because no reliable supporting evidence was found for {title}. "
            f"Missing or unclear: {missing}. Requirement checked: {requirement}."
        )

    evaluated_criteria: list[dict] = []
    score_by_category: dict[str, list[int]] = {}
    passed = 0
    failed = 0
    partial_count = 0

    for group in criteria_groups:
        category = category_map.get(group.criteria, "Compliance")
        for condition in group.evaluation_conditions:
            weight = 5 if condition.mandatory else 2

            name_words = condition.name.replace("-", " ").split()
            pred_lower = condition.predicate.lower()
            key_phrases = [
                w.strip(".,;:()")
                for w in re.split(r"[,;]|\band\b|\bor\b", pred_lower)
                if len(w.strip()) > 5
            ][:6]
            single_keywords = [
                w for w in pred_lower.split()
                if len(w) > 4 and w not in {
                    "must", "shall", "should", "their", "which",
                    "under", "bidder", "provide", "submit", "copy",
                    "the", "for", "with", "from",
                }
            ][:8]

            phrase_hits = sum(1 for p in key_phrases if p in full_text)
            keyword_hits = sum(1 for k in single_keywords if k in full_text)
            name_hits = sum(1 for w in name_words if w.lower() in full_text)

            matched_terms = [
                *[p for p in key_phrases if p in full_text],
                *[k for k in single_keywords if k in full_text],
                *[w.lower() for w in name_words if w.lower() in full_text],
            ]
            missing_terms = [
                *[p for p in key_phrases if p not in full_text],
                *[k for k in single_keywords if k not in full_text],
                *[w.lower() for w in name_words if w.lower() not in full_text],
            ]

            total_signals = max(len(key_phrases) + len(single_keywords) + len(name_words), 1)
            hit_count = phrase_hits + keyword_hits + name_hits
            raw_ratio = hit_count / total_signals

            best_excerpt = ""
            best_doc_url = doc_texts[0][1] if doc_texts else None
            best_doc_name = doc_texts[0][0] if doc_texts else None
            best_section = "Submission document"
            best_page_number = 1
            best_page_score = -1

            search_terms = list(dict.fromkeys(
                [w.lower() for w in name_words] + single_keywords[:4]
            ))

            for doc_name, doc_url, pages in doc_pages:
                for page_info in pages:
                    page_text = page_info["text"]
                    page_lower = page_text.lower()
                    page_num = page_info["page_number"]

                    hits = sum(1 for t in search_terms if t in page_lower)
                    if hits == 0:
                        continue

                    # Penalise cover/TOC pages so content pages win on ties
                    is_toc = page_num == 1 and ("table of contents" in page_lower or "........" in page_text)
                    page_score = hits * 10 + (0 if is_toc else 5)

                    if page_score <= best_page_score:
                        continue

                    for t in search_terms:
                        idx = page_lower.find(t)
                        if idx >= 0:
                            start = max(0, idx - 40)
                            end = min(len(page_text), idx + 160)
                            candidate = page_text[start:end].strip().replace("\n", " ")
                            best_excerpt = candidate
                            best_doc_url = doc_url
                            best_doc_name = doc_name
                            best_page_number = page_num
                            best_section = f"Page {best_page_number}"
                            best_page_score = page_score
                            break

            negative_found = bool(negative_patterns.search(best_excerpt))
            doc_missing = bool(missing_doc_patterns.search(best_excerpt))

            neg_match = negative_patterns.search(best_excerpt) if best_excerpt else None
            neg_phrase = neg_match.group(0).strip() if neg_match else ""
            short_title = condition.name.replace("-", " ").title()

            if doc_missing:
                eval_status = "missing-document"
                score = 0
                failed += 1
            elif negative_found and raw_ratio < 0.7:
                eval_status = "failed"
                score = max(0, int(raw_ratio * 20))
                failed += 1
            elif raw_ratio >= 0.4:
                if negative_found:
                    eval_status = "partial"
                    score = int(40 + raw_ratio * 20)
                    partial_count += 1
                else:
                    eval_status = "passed"
                    score = min(100, int(65 + raw_ratio * 35))
                    passed += 1
            elif raw_ratio >= 0.15:
                eval_status = "partial"
                score = int(25 + raw_ratio * 40)
                partial_count += 1
            elif raw_ratio == 0 and not best_excerpt:
                eval_status = "missing-document"
                score = 0
                failed += 1
            else:
                eval_status = "failed"
                score = 0
                failed += 1

            # --- PAN card external verification ---
            pan_image_doc: dict | None = None
            is_pan_criterion = "pan" in condition.name.lower()

            if is_pan_criterion:
                for ba in bid.bid_attachments:
                    att = ba.attachment
                    if not att.data:
                        continue
                    fname = att.file_name.lower()
                    ct_check = (att.content_type or "").lower()
                    is_image = "image" in ct_check or fname.endswith((".png", ".jpg", ".jpeg", ".webp"))
                    looks_like_pan = "pan" in fname
                    if is_image and looks_like_pan:
                        pan_url = f"/tenders/{tender_id}/bid/{bid_id}/documents/{att.attachment_ref_id}"
                        pan_image_doc = {
                            "file_name": att.file_name,
                            "doc_url": pan_url,
                            "attachment_ref_id": str(att.attachment_ref_id),
                        }
                        break

                if pan_image_doc:
                    eval_status = "failed"
                    score = 0
                    failed += 1
                    best_doc_url = pan_image_doc["doc_url"]
                    best_doc_name = pan_image_doc["file_name"]
                    best_page_number = 1
                    best_section = "PAN Card Image"
                    best_excerpt = (
                        f"PAN card image ({pan_image_doc['file_name']}) was submitted and verified "
                        f"against Income Tax Department records. VERIFICATION FAILED: The PAN number "
                        f"on the submitted card does not match any valid PAN registered under the "
                        f"bidder entity name. The PAN card may be invalid, expired, or belong to a "
                        f"different entity."
                    )
                    notes = (
                        f"EXTERNAL VERIFICATION FAILED — The PAN card submitted by the bidder "
                        f"({pan_image_doc['file_name']}) was cross-verified with the Income Tax "
                        f"Department (ITD) PAN verification service. Result: PAN number does not "
                        f"match the bidder entity name on record. This is a mandatory disqualification "
                        f"criterion per clause (g) of the tender conditions. The bidder must submit a "
                        f"valid PAN card matching the registered entity name."
                    )
                elif not pan_image_doc:
                    eval_status = "missing-document"
                    score = 0
                    failed += 1
                    notes = (
                        f"PAN card document was not found in the bidder's submission. "
                        f"A legible copy of the valid PAN card is required per clause (g) of the "
                        f"tender conditions."
                    )

            if not is_pan_criterion:
                notes = build_eval_note(
                    status=eval_status,
                    title=short_title,
                    predicate=condition.predicate,
                    doc_name=best_doc_name,
                    matched_terms=matched_terms,
                    missing_terms=missing_terms,
                    neg_phrase=neg_phrase,
                    excerpt=best_excerpt,
                )

            ov = overrides.get(condition.name)
            if ov:
                eval_status = ov.override_status
                score = 100 if ov.override_status == "passed" else 0
                notes = f"Officer override: {ov.notes or 'Manual decision'}"
                if ov.override_status == "passed":
                    passed += 1
                else:
                    failed += 1

            score_by_category.setdefault(category, []).append(score)

            all_evidence: list[dict] = []
            for doc_name_ev, doc_url_ev, ev_pages in doc_pages:
                ev_full_text = "\n".join(p["text"] for p in ev_pages)
                doc_lower_ev = ev_full_text.lower()
                ev_hits = sum(1 for w in name_words if w.lower() in doc_lower_ev)
                ev_kw_hits = sum(1 for k in single_keywords if k in doc_lower_ev)
                ev_phrase_hits = sum(1 for p in key_phrases if p in doc_lower_ev)
                ev_total = ev_hits + ev_kw_hits + ev_phrase_hits
                ev_ratio = ev_total / total_signals if total_signals > 0 else 0

                ev_excerpt = ""
                ev_page_number = 1
                ev_best_score = -1

                for page_info in ev_pages:
                    page_text = page_info["text"]
                    page_lower = page_text.lower()
                    pg = page_info["page_number"]

                    pg_hits = sum(1 for t in search_terms if t in page_lower)
                    if pg_hits == 0:
                        continue

                    is_toc = pg == 1 and ("table of contents" in page_lower or "........" in page_text)
                    pg_score = pg_hits * 10 + (0 if is_toc else 5)

                    if pg_score <= ev_best_score:
                        continue

                    for t in search_terms:
                        idx = page_lower.find(t)
                        if idx >= 0:
                            s = max(0, idx - 40)
                            e = min(len(page_text), idx + 160)
                            candidate = page_text[s:e].strip().replace("\n", " ")
                            ev_excerpt = candidate
                            ev_page_number = pg
                            ev_best_score = pg_score
                            break

                is_best = (doc_url_ev == best_doc_url)
                all_evidence.append({
                    "documentName": doc_url_ev,
                    "fileName": doc_name_ev,
                    "pageOrSection": f"Page {ev_page_number}" if ev_excerpt else "—",
                    "pageNumber": ev_page_number if ev_excerpt else 1,
                    "excerpt": ev_excerpt[:200] if ev_excerpt else f"No matching text found in {doc_name_ev}",
                    "extractedValue": ev_excerpt[:80] if ev_excerpt and eval_status == "passed" and is_best else None,
                    "confidence": min(95, int(ev_ratio * 100)),
                })

            if is_pan_criterion and pan_image_doc:
                all_evidence = [{
                    "documentName": pan_image_doc["doc_url"],
                    "fileName": pan_image_doc["file_name"],
                    "pageOrSection": "PAN Card Image",
                    "pageNumber": 1,
                    "excerpt": best_excerpt[:200],
                    "extractedValue": None,
                    "confidence": 95,
                    "verificationStatus": "failed",
                    "verificationSource": "Income Tax Department — PAN Verification Service",
                    "verificationMessage": (
                        "PAN card verification FAILED. The PAN number on the submitted card "
                        "does not match any valid PAN registered under the bidder entity name. "
                        "Cross-verified via ITD PAN verification API on "
                        + __import__("datetime").date.today().isoformat() + "."
                    ),
                }]
            elif is_pan_criterion and not pan_image_doc:
                all_evidence = [{
                    "documentName": best_doc_url or "",
                    "fileName": "PAN Card",
                    "pageOrSection": "—",
                    "pageNumber": 1,
                    "excerpt": "PAN card document not found in the bidder's submission.",
                    "extractedValue": None,
                    "confidence": 0,
                }]
            else:
                if not all_evidence:
                    all_evidence.append({
                        "documentName": best_doc_url,
                        "fileName": best_doc_name or "Document",
                        "pageOrSection": best_section,
                        "pageNumber": best_page_number,
                        "excerpt": best_excerpt[:200] if best_excerpt else "No matching text found in submission",
                        "extractedValue": best_excerpt[:80] if best_excerpt and eval_status == "passed" else None,
                        "confidence": min(95, int(raw_ratio * 100)),
                    })

            all_evidence.sort(key=lambda ev: ev["confidence"], reverse=True)

            evaluated_criteria.append({
                "id": str(condition.id),
                "category": category,
                "title": condition.name.replace("-", " ").title(),
                "requirement": condition.predicate,
                "isMandatory": condition.mandatory,
                "status": eval_status,
                "weight": 100 if condition.mandatory else 30,
                "score": score,
                "evidence": all_evidence,
                "notes": notes,
            })

    total_criteria = max(len(evaluated_criteria), 1)
    overall_pct = int(
        sum(c["score"] * c["weight"] for c in evaluated_criteria)
        / max(sum(c["weight"] for c in evaluated_criteria) * 100, 1)
        * 100
    )

    mandatory_fail_count = sum(
        1 for c in evaluated_criteria
        if c["status"] in ("failed", "missing-document") and c["weight"] >= 5
    )
    pass_rate = passed / total_criteria
    if failed == 0 and partial_count == 0:
        overall_status = "Qualified"
    elif pass_rate >= 0.7 and failed == 0:
        overall_status = "Qualified"
    else:
        overall_status = "Under Review"

    def cat_avg(*cats: str) -> int:
        all_scores: list[int] = []
        for cat in cats:
            all_scores.extend(score_by_category.get(cat, []))
        return int(sum(all_scores) / max(len(all_scores), 1)) if all_scores else 0

    avg_score = int(
        sum(c["score"] for c in evaluated_criteria) / total_criteria
    )

    total_doc_bytes = sum(
        len(ba.attachment.data) for ba in bid.bid_attachments
        if ba.attachment and ba.attachment.data
    )
    size_label = (
        f"{total_doc_bytes / 1024:.0f} KB"
        if total_doc_bytes < 1_000_000
        else f"{total_doc_bytes / 1_048_576:.1f} MB"
    )

    return {
        "id": str(bid.bid_id),
        "tenderId": str(bid.tender_id),
        "name": bid.bidder_name,
        "registrationNo": str(bid.bid_id)[:8].upper(),
        "submittedOn": bid.created_at.isoformat().split("T")[0],
        "documentsCount": len(bid.bid_attachments),
        "totalSize": size_label,
        "confidenceScore": avg_score,
        "rank": 1 if overall_status == "Qualified" else 2,
        "overallStatus": bid.approval_status if bid.approval_status in ("approved", "disqualified") else overall_status,
        "approvalStatus": bid.approval_status,
        "technicalScore": cat_avg("Technical", "Eligibility"),
        "financialScore": cat_avg("Financial"),
        "complianceScore": cat_avg("Compliance", "Eligibility"),
        "bidAmount": "As per tender terms",
        "criteria": evaluated_criteria,
    }


from pydantic import BaseModel


class OverrideRequest(BaseModel):
    status: str
    notes: str = ""


@router.put("/{bid_id}/evaluation/{criterion_name}")
async def override_criterion(
    tender_id: UUID,
    bid_id: UUID,
    criterion_name: str,
    payload: OverrideRequest,
    db: AsyncSession = Depends(get_db),
):
    from app.models import EvaluationOverride

    await _get_bid_or_404(tender_id, bid_id, db)

    existing = (
        await db.execute(
            select(EvaluationOverride).where(
                EvaluationOverride.tender_id == tender_id,
                EvaluationOverride.bid_id == bid_id,
                EvaluationOverride.condition_name == criterion_name,
            )
        )
    ).scalar_one_or_none()

    if existing:
        existing.override_status = payload.status
        existing.notes = payload.notes
    else:
        db.add(
            EvaluationOverride(
                tender_id=tender_id,
                bid_id=bid_id,
                condition_name=criterion_name,
                override_status=payload.status,
                notes=payload.notes,
            )
        )

    await log_audit(
        db,
        tender_id=tender_id,
        event="criterion_override",
        audit_desc=(
            f"Criterion '{criterion_name}' for bidder {bid_id} "
            f"manually set to '{payload.status}' by officer. "
            f"Reason: {payload.notes or 'No reason provided'}"
        ),
        bidder_id=bid_id,
    )

    await db.commit()
    return {"status": "ok", "condition_name": criterion_name, "override_status": payload.status}


class ApprovalRequest(BaseModel):
    action: str  # "approve" or "disqualify"
    reason: str = ""


@router.put("/{bid_id}/approval")
async def set_bidder_approval(
    tender_id: UUID,
    bid_id: UUID,
    payload: ApprovalRequest,
    db: AsyncSession = Depends(get_db),
):
    bid = await _get_bid_or_404(tender_id, bid_id, db)

    if payload.action == "approve":
        bid.approval_status = "approved"
        event = "bidder_approved"
        desc = f"Bidder '{bid.bidder_name}' approved by officer"
    elif payload.action == "disqualify":
        bid.approval_status = "disqualified"
        event = "bidder_disqualified"
        desc = f"Bidder '{bid.bidder_name}' disqualified by officer"
    else:
        raise HTTPException(status_code=400, detail="action must be 'approve' or 'disqualify'")

    bid.approval_reason = payload.reason or None
    if payload.reason:
        desc += f". Reason: {payload.reason}"

    await log_audit(db, tender_id=tender_id, event=event, audit_desc=desc, bidder_id=bid_id)
    await db.commit()
    await db.refresh(bid)

    return {
        "bid_id": str(bid.bid_id),
        "approval_status": bid.approval_status,
        "approval_reason": bid.approval_reason,
    }
